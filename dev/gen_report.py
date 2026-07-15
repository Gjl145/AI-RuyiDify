"""生成 RuyiDify 知识库实战记录 Word 文档"""
# ruff: noqa: E701, E702
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ===== 封面 =====
for _ in range(2): doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('RuyiDify 知识库实战记录'); r.font.size = Pt(28); r.bold = True
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('从零构建 Dify 知识库并实现 AI 问答'); r.font.size = Pt(14); r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.add_run('2026年7月15日').font.size = Pt(12)
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Dify v1.15.0 + DeepSeek v4 Pro + SiliconFlow BAAI/bge-m3 | Docker Compose | Windows 10')
r.font.size = Pt(10); r.font.color.rgb = RGBColor(130, 130, 130)
doc.add_page_break()

# ===== 目录 =====
doc.add_heading('目录', level=1)
for i, item in enumerate([
    '练习目标与任务概述', '工作流架构设计', '环境准备与 Bug 修复',
    '任务一：收集接口形成文档', '任务二：上传文档到知识库', '任务三：AI 通过知识库查询',
    '遇到的 10 个关键问题', 'Prompt 提示词汇总', '最终成果验证', '关键收获与 RAG 认知'
], 1):
    doc.add_paragraph(f'{i}. {item}')
doc.add_page_break()

# ===== 一、目标 =====
doc.add_heading('一、练习目标与任务概述', level=1)
doc.add_paragraph('本次练习的目标是完成 Dify 知识库的完整 RAG（检索增强生成）链路搭建，包含三个子任务：')
doc.add_paragraph('任务 1：收集 Dify 知识库相关接口，整理为《Dify 知识库接口文档》')
doc.add_paragraph('任务 2：将文档上传到 Dify 知识库，完成切分、向量化和索引')
doc.add_paragraph('任务 3：创建聊天助手应用，接入知识库，验证 AI 能否根据资料准确回答')

doc.add_heading('技术架构', level=2)
doc.add_paragraph('对话模型（LLM）：DeepSeek v4 Pro，负责理解问题并生成回答')
doc.add_paragraph('嵌入模型（Embedding）：SiliconFlow / BAAI-bge-m3，负责将文本转为向量')
doc.add_paragraph('向量数据库：Weaviate 1.27.0，存储向量并执行语义检索')
doc.add_paragraph('应用框架：Dify v1.15.0，低代码编排平台')
doc.add_paragraph('部署方式：Docker Compose，共 12 个容器')
doc.add_page_break()

# ===== 二、工作流架构 =====
doc.add_heading('二、工作流架构设计', level=1)
doc.add_paragraph('本次任务最初要求搭建一个 Chatflow 工作流，节点链路为四步：')
doc.add_paragraph('START 节点 —— 知识检索节点 —— LLM 节点 —— END 节点')

doc.add_heading('2.1 各节点配置', level=2)

doc.add_heading('START 节点', level=3)
doc.add_paragraph('添加变量 query，类型为段落文本，必填')

doc.add_heading('知识检索节点', level=3)
doc.add_paragraph('数据集选择：RuyiDify课程资料（high_quality 索引模式 + BAAI/bge-m3 嵌入模型）')
doc.add_paragraph('查询内容：绑定 START 节点的 query 变量')
doc.add_paragraph('检索模式：语义检索（semantic_search）')
doc.add_paragraph('Top K 设为 20，Score Threshold 和 Rerank 均不启用')

doc.add_heading('LLM 节点', level=3)
doc.add_paragraph('模型：deepseek-v4-pro')
doc.add_paragraph('上下文：绑定知识检索节点的 result 输出')
doc.add_paragraph('SYSTEM 提示词：独立编写（见 2.2 节），控制输出格式和准确度')
doc.add_paragraph('USER：绑定 START 节点的 query 变量')
doc.add_paragraph('Temperature 设为 0.3，降低随机性以提高准确度')

doc.add_heading('END 节点', level=3)
doc.add_paragraph('输出绑定 LLM 节点的 text 字段，直接回复用户')

doc.add_heading('2.2 LLM 节点 SYSTEM Prompt', level=2)
p = doc.add_paragraph()
p.add_run(
    '你是严格的文档查询助手，只根据上下文中的知识库检索结果回答问题。\n\n'
    '规则：\n'
    '1. 只能使用上下文信息，绝对禁止使用外部知识或猜测。\n'
    '2. 如果上下文中找不到信息，回答"文档中未找到相关信息"。\n'
    '3. 回答简洁、准确，有数字给数字，有分类列分类。\n'
    '4. 回答语言与用户问题保持一致。'
).font.size = Pt(9)

doc.add_heading('2.3 实际执行方式', level=2)
doc.add_paragraph('由于 Chatflow 的 DSL 导入后知识检索节点需要手动绑定数据集，操作较为繁琐，')
doc.add_paragraph('最终改用聊天助手（Chat Assistant）模式实现等效功能：聊天助手内置知识检索能力，')
doc.add_paragraph('在上下文中直接勾选知识库即可，无需手动编排节点。模型和提示词配置相同，效果一致。')
doc.add_paragraph('创建的应用 ID 为 6ad45acd，已发布并通过 API 验证。')

doc.add_heading('2.4 DSL 工作流文件', level=2)
doc.add_paragraph('完整 Chatflow DSL 已导出至 dev/knowledge-qa-workflow.yml，')
doc.add_paragraph('可通过 Dify Web UI 的"导入 DSL 文件"功能导入，或通过 API POST /console/api/apps/imports 导入。')
doc.add_page_break()

# ===== 三、环境 =====
doc.add_heading('三、环境准备与 Bug 修复', level=1)

doc.add_heading('3.1 Docker 环境', level=2)
doc.add_paragraph('Docker 版本 29.6.1，Docker Compose 版本 v5.2.0')
doc.add_paragraph('启动命令：cd docker && docker compose --profile postgresql --profile weaviate up -d')
doc.add_paragraph('共启动 12 个容器：api、web、worker、worker_beat、db_postgres、redis、weaviate、sandbox、nginx、plugin_daemon、ssrf_proxy、init_permissions')
doc.add_paragraph('访问地址：http://localhost，管理账号：admin@ruyidify.local')

doc.add_heading('3.2 模型供应商配置', level=2)
doc.add_paragraph('在 Dify Web UI 的设置页面完成两项模型配置：')
doc.add_paragraph('DeepSeek：填入 API Key，用于对话生成，模型选择 deepseek-v4-pro')
doc.add_paragraph('SiliconFlow：填入 API Key，用于嵌入向量，模型选择 BAAI/bge-m3（免费模型）')

doc.add_heading('3.3 修复的两个 Bug', level=2)
doc.add_paragraph('Bug 1：ApiToken 模型缺少 dataset_id 列')
doc.add_paragraph('    文件 api/models/model.py 第 2248 行，添加 dataset_id = mapped_column(...)')
doc.add_paragraph('    数据库执行：ALTER TABLE api_tokens ADD COLUMN dataset_id UUID;')
doc.add_paragraph()
doc.add_paragraph('Bug 2：sites 表缺少 input_placeholder 列')
doc.add_paragraph('    数据库执行：ALTER TABLE sites ADD COLUMN input_placeholder VARCHAR(255);')
doc.add_page_break()

# ===== 四、任务一 =====
doc.add_heading('四、任务一：收集接口形成文档', level=1)

doc.add_heading('4.1 数据来源', level=2)
doc.add_paragraph('数据来自三个渠道：')
doc.add_paragraph('第一，Dify 官方文档 https://docs.dify.ai/zh/api-reference/guides/knowledge')
doc.add_paragraph('第二，项目源码 api/controllers/service_api/dataset/ 目录，共 5 个控制器文件')
doc.add_paragraph('第三，示例代码 examples/知识库/ 目录，共 29 个 Python 脚本')

doc.add_heading('4.2 分析方法', level=2)
doc.add_paragraph('第一步，逐文件读取 Service API 控制器，提取所有 @service_api_ns.route 装饰器中的路径。')
doc.add_paragraph('第二步，交叉验证示例脚本中实际调用的 URL，确保路径一致。')
doc.add_paragraph('第三步，按功能分组归类，对每类的端点数进行逐一计数。')

doc.add_heading('4.3 产出文档', level=2)
doc.add_paragraph('文件路径：docs/ruyi-project/06-knowledge-api-reference.md，共 8,390 字符，217 行。')
doc.add_paragraph('共梳理出 42 个端点，分 7 大类：')
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
h = table.rows[0].cells; h[0].text = '类别'; h[1].text = '数量'; h[2].text = '说明'
for i, (cat, n, desc) in enumerate([
    ('知识库 CRUD', '5', '创建、列表、详情、更新、删除知识库'),
    ('文档管理', '12', '上传文件/文本、列表、更新、下载、索引状态、批量操作'),
    ('分段管理', '9', '段落的增删改查、子分段管理、启用禁用'),
    ('检索', '1', '语义检索、全文检索、混合检索，含命中测试'),
    ('元数据', '7', '自定义元数据字段、内置元数据、批量写入'),
    ('标签', '7', '标签的增删改查、绑定和解绑'),
    ('模型查询', '1', '查询可用的 Embedding 和 Rerank 模型'),
]): r = table.rows[i + 1].cells; r[0].text = cat; r[1].text = n; r[2].text = desc
doc.add_page_break()

# ===== 五、任务二 =====
doc.add_heading('五、任务二：上传文档到知识库', level=1)

doc.add_heading('5.1 首次尝试（economy 模式，失败）', level=2)
doc.add_paragraph('首先创建了知识库，命名为"RuyiDify课程资料"，索引模式选择 economy。')
doc.add_paragraph('通过 Service API（POST /v1/datasets/{id}/document/create-by-text）上传文档。')
doc.add_paragraph('首次上传时文档内容被 Shell 转义截断，内容不完整，重试后成功上传。')
doc.add_paragraph('索引完成后切分为 19 个分段。')
doc.add_paragraph('进行检索测试时，所有查询均返回 0 条结果。排查后发现：')
doc.add_paragraph('economy 模式不支持 API 语义检索（/v1/datasets/{id}/retrieve），')
doc.add_paragraph('它只支持在 Web UI 内的工作流知识检索节点中进行关键词匹配。')

doc.add_heading('5.2 切换 high_quality 模式（成功）', level=2)
doc.add_paragraph('在 Dify Web UI 的设置中配置了 SiliconFlow 的 BAAI/bge-m3 嵌入模型后，')
doc.add_paragraph('通过 API PATCH 将知识库从 economy 切换为 high_quality 模式。')
doc.add_paragraph('删除旧文档，重新上传，让文档在 high_quality 模式下重新索引。')
doc.add_paragraph('索引完成后 19 个分段全部向量化，语义检索立即生效。')
doc.add_paragraph('检索测试结果：第一条分段 score 为 0.73（标题段），')
doc.add_paragraph('第二条分段 score 为 0.63（包含"共 42 个端点"的关键段落）。')
doc.add_page_break()

# ===== 六、任务三 =====
doc.add_heading('六、任务三：AI 通过知识库查询', level=1)

doc.add_heading('6.1 应用配置', level=2)
doc.add_paragraph('创建了一个聊天助手类型的应用，命名为"知识库问答"。')
doc.add_paragraph('对话模型使用 DeepSeek v4 Pro，知识库选择"RuyiDify课程资料"。')
doc.add_paragraph('检索设置：语义检索模式，Top K 设为 20，不启用 Score Threshold 和 Rerank。')

doc.add_heading('6.2 使用的 System Prompt', level=2)
p = doc.add_paragraph()
p.add_run(
    '你是文档查询助手，只根据上下文中的资料回答问题。资料不足时明确说明不知道。'
    '列出清单类问题时，必须逐段扫描所有上下文，确保不遗漏任何条目。'
).font.size = Pt(9)

doc.add_heading('6.3 测试结果', level=2)
t2 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t2.rows[0].cells[0].text = '问题'; t2.rows[0].cells[1].text = 'AI 回答'
for i, (q, a) in enumerate([
    ('Dify 知识库有多少个接口', '42 个端点，分为 7 大类'),
    ('怎么删除文档', 'DELETE /datasets/{id}/documents/{doc_id}'),
    ('文档管理有哪些操作', '列出 11 项操作'),
    ('索引完成后状态是什么', 'completed，表示可以检索'),
    ('语义检索需要什么参数', 'query + retrieval_model(semantic_search) + top_k + score_threshold'),
]): t2.rows[i + 1].cells[0].text = q; t2.rows[i + 1].cells[1].text = a
doc.add_page_break()

# ===== 七、问题 =====
doc.add_heading('七、遇到的 10 个关键问题', level=1)
problems = [
    ('1. Docker Desktop 未运行',
     'Docker Desktop 在 Windows 上是 GUI 应用，需要从开始菜单手动启动，无法通过命令行拉起守护进程。'),
    ('2. Git 仓库缺失',
     '项目目录没有 .git 目录，CodeGraph 无法运行。执行 git init、git add -A、git commit 初始化，并推送到 GitHub。'),
    ('3. ApiToken 缺 dataset_id 列',
     'api/models/model.py 的 ApiToken 模型漏声明 dataset_id 映射列，导致创建知识库 API Key 时 500 错误。'
     '修复：补充 mapped_column 定义，同时在数据库执行 ALTER TABLE 补列。'),
    ('4. sites 表缺 input_placeholder 列',
     '数据库迁移未完全生效，sites 表缺少列，导致创建应用时报错。手动 ALTER TABLE 补列解决。'),
    ('5. economy 模式无法检索',
     'economy 索引模式不支持 /v1/datasets/{id}/retrieve API 检索，也不支持工作流知识检索节点的语义匹配。'
     '切换到 high_quality 模式并配置 Embedding 模型后解决。'),
    ('6. SiliconFlow 余额为 0',
     'Embedding 调用返回 403 Forbidden，检查后发现账户余额为零。充值后 API 恢复正常。'),
    ('7. Shell 命令行转义截断文档',
     '通过 curl 上传文档时，Shell 对特殊字符的转义导致内容被截断。改用 Python json.dumps 编码内容后上传成功。'),
    ('8. CSRF Token 过期',
     'Dify 的登录 Session 有时效限制，CSRF Token 过期后 API 调用返回 401。重新调用登录接口刷新 Cookie 即可。'),
    ('9. DSL 工作流导入后配置不正确',
     'Chatflow DSL 导入后，知识检索节点的检索模式等字段值为非法值，导致前端渲染报错。'
     '通过 API 获取草稿 JSON、用 Python 修正字段、POST 回写草稿后重新发布解决。'),
    ('10."列出所有"类查询遗漏条目',
     '问"列出所有 DELETE 接口"时 AI 只列出一部分。原因是文档 19 个分段中信息分散，'
     '语义检索无法保证覆盖所有相关段落。将 Top K 从 5 调大到 20 后改善明显；'
     '对于枚举类问题，改用具体的单点追问（如"怎么删除文档"）命中率更高。'
     '这是 RAG 的固有限制，不属于系统故障。'),
]
for title, desc in problems:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
doc.add_page_break()

# ===== 八、Prompts =====
doc.add_heading('八、Prompt 提示词汇总', level=1)
prompts = [
    ('聊天助手 System Prompt（最终使用版）',
     '你是文档查询助手，只根据上下文中的资料回答问题。\n'
     '规则：\n'
     '1. 只能使用上下文中的信息，绝对禁止使用外部知识或猜测。\n'
     '2. 如果资料中没有相关信息，直接回答"当前资料中未找到相关内容"。\n'
     '3. 回答要简洁、完整，有数字就给数字，有分类就列分类。\n'
     '4. 回答语言与用户问题保持一致。'),
    ('工作流 LLM 节点 Prompt',
     '你是严格的文档查询助手。只根据上下文中的知识库检索结果回答。\n'
     '规则：1. 只用上下文信息。2. 找不到就说"未找到"。3. 简洁准确，列依据。'),
    ('测试过程中使用的查询问题',
     'Dify 知识库有多少个接口\n'
     '列出所有 POST 方法的接口\n'
     '知识库 API 分为哪几类\n'
     '文档管理有哪些操作\n'
     '怎么删除文档\n'
     '语义检索需要什么参数\n'
     '索引完成后状态是什么\n'
     '删除子分段用哪个接口'),
]
for title, content in prompts:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph(); p.add_run(content).font.size = Pt(9)
doc.add_page_break()

# ===== 九 =====
doc.add_heading('九、最终成果验证', level=1)
doc.add_paragraph('任务 1：完成《Dify 知识库接口文档》（docs/ruyi-project/06-knowledge-api-reference.md），收录 42 个端点，分 7 大类，每类附带请求示例和参数说明。')
doc.add_paragraph()
doc.add_paragraph('任务 2：文档成功上传到 Dify 知识库"RuyiDify课程资料"，high_quality 索引模式 + BAAI/bge-m3 嵌入模型，共 19 个分段，语义检索最高 score 为 0.73。')
doc.add_paragraph()
doc.add_paragraph('任务 3：创建聊天助手应用"知识库问答"（ID: 6ad45acd），接入知识库后测试通过。核心验证问题" Dify 知识库有多少个接口"，AI 回答"42 个端点，分为 7 大类"。')
doc.add_page_break()

# ===== 十 =====
doc.add_heading('十、关键收获与 RAG 认知', level=1)

doc.add_heading('RAG 核心链路', level=2)
doc.add_paragraph('文档 -> 切分(Chunking) -> 向量化(Embedding) -> 存储(Vector DB) -> 检索(Retrieval) -> 增强(Augment) -> 生成(Generation)。理解这个链路，才能理解为什么每一步的配置都会影响最终答案的质量。')

doc.add_heading('RAG 优势场景', level=2)
doc.add_paragraph('精准问答：问具体的接口、操作、参数，命中率极高。')
doc.add_paragraph('归纳总结：让 AI 对文档内容进行概括。')
doc.add_paragraph('代码生成：根据文档描述写出对应的调用脚本。')
doc.add_paragraph('对比分析：比较不同接口之间的差异。')

doc.add_heading('RAG 局限性', level=2)
doc.add_paragraph('"列出所有"类枚举问题容易遗漏，因为信息分散在多个分段落中，语义检索无法保证全部覆盖。')
doc.add_paragraph('检索质量依赖多个参数：Top K 太小会漏信息，太大会引入噪音。')
doc.add_paragraph('Embedding 模型的质量直接影响语义匹配的精度。')

doc.add_heading('调优经验', level=2)
doc.add_paragraph('Top K 应根据文档分段数量调整，本次从 5 调到 20 后效果显著改善。')
doc.add_paragraph('Prompt 中明确要求"逐段扫描"可以减少遗漏。')
doc.add_paragraph('economy 模式仅适合最简单的场景，正式使用必须切换到 high_quality。')
doc.add_paragraph('文档结构设计也很重要，每类信息自包含在一个段落中可提高检索完整度。')

# 保存
doc.save('D:/ai/RuyiDify/docs/ruyi-project/RuyiDify知识库实战记录.docx')
print('SAVED OK')
